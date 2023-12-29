import streamlit as st
import pyperclip
import os
from docx import Document
import subprocess
import sys
import re
import json

# 读取 JSON 文件
st.set_page_config(layout="wide")
current_dir = os.getcwd()  # 获取当前目录
parent_dir = os.path.dirname(current_dir)  # 获取上一级目录
config_file_path = os.path.normpath('./config.json')
basic_data = []
with open(config_file_path, 'r', encoding='utf-8') as file:
    basic_data = json.load(file)
def replace_path_variables(path, current_dir, parent_dir):
    path = path.replace('current_dir + ', current_dir)
    path = path.replace('current_dir', current_dir)
    path = path.replace('parent_dir + ', parent_dir)
    path = path.replace('parent_dir', parent_dir)
    return path
for item in basic_data:
    for key, value in basic_data[item].items():
        if key.startswith("doc_file") or key == "ignore_folder":
            if key == "ignore_folder" and value == "":
                basic_data[item][key] = os.path.dirname(basic_data[item]["doc_file1"])
            else:
                basic_data[item][key] = replace_path_variables(value, current_dir, parent_dir)
# print('-----------------')
# print('basic_data: ', basic_data)


# 定义基础数据
ignore_folder = current_dir
doc_file1 = ""
doc_file2 = ""
doc_files = []

# 显示标题
col1, col2 = st.columns([5,1])
with col1:
    st.title('英语试题查重工具')
with col2:
    st.markdown(" ", unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    st.text("Designed by xie")


# 获得关键字列表
key_words_list = [basic_data[item]["key_words"] for item in basic_data]
key_words = st.selectbox('输入搜索关键字', key_words_list)
# print('key_words_list: ', key_words_list)


# 得到题库地址
for key, value in basic_data.items():
    if value.get('key_words') == key_words:
        st.info(f'学生习题搜索关键字：{value.get("key_words")}')
        for file_key, file_path in value.items():
            if file_key.startswith('doc_file'):
                doc_files.append(file_path)
st.warning("题库地址:")
for i, file in enumerate(doc_files):
    globals()[f'doc_file{i + 1}'] = file
    path_parts = os.path.splitdrive(file)[1].split(os.sep)[-3:]
    path = os.path.join(*path_parts)
    exec(f'st.text("{path}")')


# 得到忽视文件夹
for item in basic_data:
    if key_words == basic_data[item]["key_words"]:
        ignore_folder = basic_data[item]["ignore_folder"]


# 显示试题
st.markdown("---")
col1, col2 = st.columns([6, 1])
with col1:
    user_input = st.text_input("输入文本",key="unique_key_1")
with col2:
    selected_language = st.selectbox("请选择一个搜索内容", ("英文", "中文"), format_func=lambda x: x)


# 输入试题
def get_source(selected_language="英文"):
    source_doc = pyperclip.paste()
    len_source = len(source_doc)
    if selected_language == "英文":
        match = re.search(r'\b\w+\b(?: \w+){7}', source_doc)
    if selected_language == "中文":
        match = re.search(r'[\u4e00-\u9fff]{15}', source_doc)

    if match:
        return match.group()
    else:
        return "匹配失败"


# 显示关键字和sidebar
key_source = get_source(selected_language)
st.sidebar.write(f'{user_input}')

if len(user_input) == 0:
    pass
elif key_source == "匹配失败":
    st.info("匹配失败，无法查询")
else:
    st.warning(f'试题搜索关键字为：{key_source}')
    st.markdown("---")
    st.text("")
    # 查找学生文件
    def get_file_path(key_words, ignore_folder=current_dir):
        file_path = []
        for root, dirs, files in os.walk(parent_dir):
            if root == ignore_folder:   # 跳过当前文件夹
                continue
            for file_name in files:
                if key_words in file_name and "~$" not in file_name:
                    file_path.append(os.path.join(root, file_name))

        return file_path
    file_paths = get_file_path(key_words, ignore_folder)

    # 检查俩个题库是否有重复
    def is_string_in_docx(source_doc, doc_files):
        result_list = []
        for doc_file in doc_files:
            doc = Document(doc_file)
            doc_text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
            if source_doc in doc_text:
                count = doc_text.count(source_doc)
                result_list.append(('docx中重复了{}次'.format(count), doc_file))

            else:
                result_list.append(('docx没有重复', doc_file))
        return result_list
    doc_list = is_string_in_docx(key_source, doc_files)

    # 检查每个学生是否重复
    def check_each(source_doc, file_paths):
        result_list = []
        for file in file_paths:
            doc = Document(file)
            doc_text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
            if source_doc in doc_text:
                count = doc_text.count(source_doc)
                result_list.append(('docx中重复了{}次'.format(count), file))
            else:
                result_list.append(('可以新加的学生', file))

        return result_list
    each_list = check_each(key_source, file_paths)


    # 显示计算重复结果
    def open_docx_file(file_path):
        if sys.platform == "win32":
            subprocess.run(["start", file_path], shell=True)
        elif sys.platform == "darwin":
            subprocess.run(["open", file_path])
        elif sys.platform == "linux":
            subprocess.run(["xdg-open", file_path])
        else:
            print("无法在当前平台上执行打开操作")
    if user_input != "":
        i = 0
        for item in doc_list:
            col1, col2 = st.columns([6, 1])
            i += 1
            with col1:
                if "没有重复" in item[0]:
                    st.success(item[0] + "： " + "/".join(item[1].split("/")[-3:]))
                else:
                    st.error(item[0] + "： " + "/".join(item[1].split("/")[-3:]))
            with col2:
                st.markdown(" ", unsafe_allow_html=True)
                if st.button(f'打开题库{i}'):
                    open_docx_file(item[1])
        st.markdown("---")
        st.text("")

        for i, item in enumerate(each_list):
            col1, col2 = st.columns([6, 1])
            with col1:
                if "可以新加的学生" in item[0]:
                    st.success(item[0] + "： " + "/".join(item[1].split("/")[-3:]))
                else:
                    st.error(item[0] + "： " + "/".join(item[1].split("/")[-3:]))
            with col2:
                if "可以新加的学生" in item[0]: 
                    st.markdown(" ", unsafe_allow_html=True)
                    if st.button(f'点击打开{i + 1}'):
                        open_docx_file(item[1])